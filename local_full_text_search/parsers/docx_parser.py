from __future__ import annotations

from pathlib import Path
from typing import Iterable

from local_full_text_search.core.errors import ParserDependencyError
from local_full_text_search.core.task_manager import CancelToken
from local_full_text_search.models.content_block import ContentBlock
from local_full_text_search.parsers.base_parser import BaseParser


class DocxParser(BaseParser):
    name = "docx"

    def supports(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".docx"

    def parse(self, file_path: Path, cancel_token: CancelToken) -> Iterable[ContentBlock]:
        try:
            from docx import Document
        except ImportError as exc:
            raise ParserDependencyError("未安装 python-docx，无法解析 DOCX") from exc
        document = Document(str(file_path))
        block_index = 0
        paragraph_no = 0
        for paragraph in document.paragraphs:
            cancel_token.wait_if_paused()
            cancel_token.throw_if_cancelled()
            paragraph_no += 1
            text = paragraph.text.strip()
            if not text:
                continue
            yield self.make_block(
                file_path,
                block_index,
                "docx_paragraph",
                f"正文第 {paragraph_no} 段",
                text,
            )
            block_index += 1
        for table_index, table in enumerate(document.tables, start=1):
            for row_index, row in enumerate(table.rows, start=1):
                cancel_token.wait_if_paused()
                cancel_token.throw_if_cancelled()
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if not cells:
                    continue
                yield self.make_block(
                    file_path,
                    block_index,
                    "docx_table_row",
                    f"表格 {table_index}，第 {row_index} 行",
                    " | ".join(cells),
                    extra={"table_index": table_index, "row_index": row_index},
                )
                block_index += 1
        for section_index, section in enumerate(document.sections, start=1):
            for part_name, container in (("页眉", section.header), ("页脚", section.footer)):
                texts = [p.text.strip() for p in container.paragraphs if p.text.strip()]
                if texts:
                    yield self.make_block(
                        file_path,
                        block_index,
                        "docx_header_footer",
                        f"{part_name} {section_index}",
                        "\n".join(texts),
                        extra={"section_index": section_index, "part": part_name},
                    )
                    block_index += 1
