from __future__ import annotations

import tempfile
import threading
import time
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document

from local_full_text_search.config.defaults import AppSettings
from local_full_text_search.core.database import DatabaseManager
from local_full_text_search.core.errors import PauseRequestedError
from local_full_text_search.core.index_manager import IndexManager, lane_for
from local_full_text_search.core.task_manager import CancelToken
from local_full_text_search.parsers.legacy_office_parser import (
    LegacyOfficeParser,
    OfficeConversionSession,
    _shared_fallback_lock,
)


class LegacyOfficeParserTests(unittest.TestCase):
    def test_u0_03r_legacy_pause_closes_owned_office_session(
        self,
    ) -> None:
        class PausingConvertedParser:
            last_status = "success"
            last_error_code = None
            last_error_message = None

            def parse(
                self,
                _path: Path,
                _token: CancelToken,
            ) -> object:
                raise PauseRequestedError("pause")

        with tempfile.TemporaryDirectory() as tmp:
            base = Path(tmp)
            source = base / "old.doc"
            source.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1")
            converted = base / "old.docx"
            converted.write_bytes(b"converted")
            parser = LegacyOfficeParser()

            with (
                patch.object(
                    parser,
                    "_cached_conversion",
                    return_value=converted,
                ),
                patch.object(
                    parser,
                    "_parser_for_converted",
                    return_value=PausingConvertedParser(),
                ),
                patch.object(parser._office_session, "close") as close,
            ):
                with self.assertRaises(PauseRequestedError):
                    list(parser.parse(source, CancelToken()))

            close.assert_called_once_with()

    def test_u0_03r_legacy_resume_skips_confirmed_converted_blocks(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            base = Path(tmp)
            source = base / "old.doc"
            source.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1")
            converted = base / "old.docx"
            document = Document()
            document.add_paragraph("LEGACY_RESUME_BLOCK_1")
            document.add_paragraph("LEGACY_RESUME_BLOCK_2")
            document.save(converted)
            parser = LegacyOfficeParser()
            progress: list[dict[str, object]] = []
            parser.configure_runtime(
                resume_cursor=1,
                progress_callback=progress.append,
            )

            with patch.object(
                parser,
                "_cached_conversion",
                return_value=converted,
            ):
                blocks = list(parser.parse(source, CancelToken()))

            text = "\n".join(block.raw_text for block in blocks)
            self.assertTrue(parser.supports_resume)
            self.assertNotIn("LEGACY_RESUME_BLOCK_1", text)
            self.assertIn("LEGACY_RESUME_BLOCK_2", text)
            self.assertEqual(
                [
                    item["cursor"]
                    for item in progress
                    if item["phase"] == "legacy_converted_block"
                ],
                [2],
            )

    def test_invalid_legacy_document_blocks_complete_index(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            base = Path(tmp)
            root = base / "files"
            root.mkdir()
            (root / "old.doc").write_bytes(b"legacy placeholder")
            db = DatabaseManager(base / "index.db")
            db.initialize()
            root_id = db.add_root(root)
            summary = IndexManager(db, AppSettings()).index_root(root_id)
            self.assertEqual(summary.failed, 1)
            diagnostics = db.failed_files()
            self.assertEqual(diagnostics[0]["parse_status"], "failed")
            self.assertEqual(diagnostics[0]["parse_error_code"], "LEGACY_INVALID_FORMAT")
            self.assertIn("progress_phase", diagnostics[0].keys())
            self.assertIn("progress_cursor", diagnostics[0].keys())
            self.assertTrue(str(diagnostics[0]["recovery_advice"]))
            self.assertFalse(db.index_readiness()["ready"])

    def test_wps_is_used_when_microsoft_word_is_unavailable(self) -> None:
        class FakeDocument:
            def SaveAs2(self, target: str, FileFormat: int) -> None:
                self.target = target
                self.file_format = FileFormat
                Path(target).write_bytes(b"converted")

            def Close(self, _save: bool) -> None:
                return

        class FakeDocuments:
            def Open(self, _source: str, ReadOnly: bool) -> FakeDocument:
                self.read_only = ReadOnly
                return FakeDocument()

        class FakeApp:
            Documents = FakeDocuments()

            def Quit(self) -> None:
                return

        with tempfile.TemporaryDirectory() as tmp:
            base = Path(tmp)
            source = base / "old.doc"
            target = base / "old.docx"
            source.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1")
            session = OfficeConversionSession()
            attempts: list[str] = []

            def application(_kind: str, prog_id: str) -> object:
                attempts.append(prog_id)
                if prog_id == "Word.Application":
                    raise RuntimeError("class not registered")
                return FakeApp()

            with patch.object(session, "_application", side_effect=application):
                result = session.convert(source, target)

            self.assertEqual(attempts, ["Word.Application", "KWPS.Application"])
            self.assertEqual(result.path, target)
            self.assertTrue(target.is_file())

    def test_legacy_formats_use_independent_application_lanes(self) -> None:
        settings = AppSettings()

        self.assertEqual(lane_for(Path("old.doc"), settings), "legacy_word")
        self.assertEqual(lane_for(Path("old.xls"), settings), "legacy_excel")
        self.assertEqual(
            lane_for(Path("old.ppt"), settings),
            "legacy_powerpoint",
        )

    def test_conversion_cache_reuses_precomputed_source_digest(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "old.doc"
            source.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1")
            parser = LegacyOfficeParser()
            parser.configure_runtime(content_digest="sha256:precomputed")

            with patch.object(
                Path,
                "open",
                side_effect=AssertionError("legacy source rehashed"),
            ):
                cache_path = parser._cache_path(source)

            self.assertTrue(cache_path.stem.startswith("precomputed.v4"))

    def test_shared_fallback_converter_is_serialized_across_lanes(self) -> None:
        barrier = threading.Barrier(2)
        state_lock = threading.Lock()
        active = 0
        maximum_active = 0

        def convert() -> None:
            nonlocal active, maximum_active
            barrier.wait()
            with _shared_fallback_lock():
                with state_lock:
                    active += 1
                    maximum_active = max(maximum_active, active)
                time.sleep(0.05)
                with state_lock:
                    active -= 1

        threads = [threading.Thread(target=convert) for _ in range(2)]
        for thread in threads:
            thread.start()
        for thread in threads:
            thread.join(timeout=2)

        self.assertTrue(all(not thread.is_alive() for thread in threads))
        self.assertEqual(maximum_active, 1)


if __name__ == "__main__":
    unittest.main()
